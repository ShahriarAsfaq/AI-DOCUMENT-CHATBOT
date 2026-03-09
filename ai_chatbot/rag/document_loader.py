"""Document loader for PDF and DOCX files with OCR fallback.

Supports extracting text from PDF and DOCX documents with metadata.
Uses PyMuPDF for primary PDF parsing and OCR fallback for scanned documents.
"""

import logging
import os
from pathlib import Path
from typing import List, Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from .document import Document

logger = logging.getLogger(__name__)


def load_document(file_path: str) -> List[Document]:
    """Load and extract text from PDF or DOCX files with OCR fallback.

    Args:
        file_path: Path to the PDF or DOCX file.

    Returns:
        List of Document objects with extracted text and metadata.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file format is not supported (not PDF or DOCX).
        Exception: For other document processing errors.
    """
    file_path = Path(file_path)

    # Validate file exists
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    # Get file extension and name
    file_ext = file_path.suffix.lower()
    file_name = file_path.name

    # Route to appropriate loader
    if file_ext == ".pdf":
        return _load_pdf_with_ocr_fallback(file_path, file_name)
    elif file_ext == ".docx":
        return _load_docx(file_path, file_name)
    else:
        raise ValueError(
            f"Unsupported file format: {file_ext}. Supported formats: .pdf, .docx"
        )


def _load_pdf_with_ocr_fallback(file_path: Path, file_name: str) -> List[Document]:
    """Extract text from PDF file with OCR fallback for scanned documents.

    Args:
        file_path: Path to the PDF file.
        file_name: Name of the file for metadata.

    Returns:
        List of Document objects, one per page.

    Raises:
        Exception: If PDF processing fails.
    """
    documents: List[Document] = []

    try:
        logger.info(f"Loading PDF: {file_name}")

        # First try PyMuPDF extraction
        pdf_docs = _extract_text_with_pymupdf(file_path, file_name)

        # Check if we got meaningful text
        total_text_length = sum(len(doc.page_content.strip()) for doc in pdf_docs)
        logger.info(f"PyMuPDF extracted {total_text_length} characters from {len(pdf_docs)} pages")

        # If text is too short, try OCR
        if total_text_length < 100:  # Very low threshold for meaningful text
            logger.warning(f"Low text content detected ({total_text_length} chars). Attempting OCR...")
            pdf_docs = _extract_text_with_ocr(file_path, file_name)
            logger.info(f"OCR extracted {sum(len(doc.page_content.strip()) for doc in pdf_docs)} characters")

        # Filter out empty pages and clean text
        valid_docs = []
        for doc in pdf_docs:
            cleaned_text = _clean_text(doc.page_content)
            if len(cleaned_text.strip()) >= 10:  # Minimum meaningful content
                valid_docs.append(Document(
                    page_content=cleaned_text,
                    metadata=doc.metadata
                ))

        if not valid_docs:
            raise ValueError(f"No meaningful text content found in PDF: {file_name}")

        logger.info(f"Successfully processed PDF with {len(valid_docs)} valid pages")
        return valid_docs

    except Exception as e:
        raise Exception(f"Error processing PDF '{file_name}': {str(e)}") from e


def _extract_text_with_pymupdf(file_path: Path, file_name: str) -> List[Document]:
    """Extract text using PyMuPDF (fitz).

    Args:
        file_path: Path to PDF file.
        file_name: Name of the file.

    Returns:
        List of Document objects.
    """
    documents = []

    try:
        doc = fitz.open(str(file_path))
        logger.info(f"Opened PDF with {len(doc)} pages")

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()

            documents.append(Document(
                page_content=text,
                metadata={
                    "page": page_num + 1,  # 1-indexed
                    "source": file_name,
                    "extraction_method": "pymupdf"
                }
            ))

        doc.close()
        return documents

    except Exception as e:
        logger.error(f"PyMuPDF extraction failed: {str(e)}")
        raise


def _extract_text_with_ocr(file_path: Path, file_name: str) -> List[Document]:
    """Extract text using OCR for scanned PDFs.

    Args:
        file_path: Path to PDF file.
        file_name: Name of the file.

    Returns:
        List of Document objects.
    """
    documents = []

    try:
        import pytesseract
        from pdf2image import convert_from_path
        import tempfile

        logger.info("Converting PDF to images for OCR...")

        # Convert PDF to images
        with tempfile.TemporaryDirectory() as temp_dir:
            images = convert_from_path(str(file_path), dpi=300, output_folder=temp_dir)

            for page_num, image in enumerate(images):
                # Perform OCR on the image
                text = pytesseract.image_to_string(image)

                documents.append(Document(
                    page_content=text,
                    metadata={
                        "page": page_num + 1,  # 1-indexed
                        "source": file_name,
                        "extraction_method": "ocr"
                    }
                ))

        logger.info(f"OCR processed {len(documents)} pages")
        return documents

    except ImportError as e:
        logger.error(f"OCR dependencies not available: {str(e)}")
        logger.error("Install pytesseract, pdf2image, and Pillow for OCR support")
        raise
    except Exception as e:
        logger.error(f"OCR extraction failed: {str(e)}")
        raise


def _clean_text(text: str) -> str:
    """Clean and normalize extracted text.

    Args:
        text: Raw extracted text.

    Returns:
        Cleaned text.
    """
    if not text:
        return ""

    # Remove excessive whitespace
    import re
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Multiple newlines to double
    text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space

    # Strip leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(line for line in lines if line)  # Remove empty lines

    return text.strip()


def _load_docx(file_path: Path, file_name: str) -> List[Document]:
    """Extract text from DOCX file.

    Args:
        file_path: Path to the DOCX file.
        file_name: Name of the file for metadata.

    Returns:
        List of Document objects. For DOCX, typically one document with all content.

    Raises:
        Exception: If DOCX processing fails.
    """
    documents = []

    try:
        docx = DocxDocument(str(file_path))
        text_content = []

        # Extract text from all paragraphs
        for para in docx.paragraphs:
            if para.text.strip():
                text_content.append(para.text)

        full_text = "\n".join(text_content)
        cleaned_text = _clean_text(full_text)

        if not cleaned_text.strip():
            raise ValueError(f"No text content found in DOCX: {file_name}")

        # Create single document for DOCX (page concept doesn't apply)
        doc = Document(
            page_content=cleaned_text,
            metadata={
                "page": 1,  # DOCX treated as single document
                "source": file_name,
                "extraction_method": "docx"
            },
        )
        documents.append(doc)

        logger.info(f"Successfully processed DOCX with {len(cleaned_text)} characters")
        return documents

    except Exception as e:
        raise Exception(f"Error processing DOCX '{file_name}': {str(e)}") from e
